#!/usr/bin/env python3
"""
generate_strobe.py
==================
Generate a STROBE checklist (Word document) for the Meditron ICU evaluation study.

STROBE = Strengthening the Reporting of Observational Studies in Epidemiology.
Cross-sectional study variant.

Output:
- output/STROBE_checklist.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).parent.parent / 'output'

# STROBE checklist items for cross-sectional studies
# (Item number, Section, Recommendation, Where addressed in manuscript)
STROBE_ITEMS = [
    # Title and abstract
    ("1", "Title and abstract", "(a) Indicate the study's design with a commonly used term in the title or the abstract",
     "Abstract: 'multi-rater assessment study' in title"),
    ("", "", "(b) Provide in the abstract an informative and balanced summary of what was done and what was found",
     "Abstract: structured with Background, Methods, Results, Conclusions"),

    # Introduction
    ("2", "Background/rationale", "Explain the scientific background and rationale for the investigation being reported",
     "Background, paragraphs 1-4"),
    ("3", "Objectives", "State specific objectives, including any prespecified hypotheses",
     "Background, final paragraph: aim stated"),

    # Methods
    ("4", "Study design", "Present key elements of study design early in the paper",
     "Methods, 'Study design and setting': subcohort analysis within MOOVE-CHUV"),
    ("5", "Setting", "Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection",
     "Methods, 'Study design and setting': Swiss tertiary university hospital; data collection dates in Data Dictionary (Jan-Mar 2025)"),
    ("6", "Participants", "Give the eligibility criteria, and the sources and methods of selection of participants",
     "Methods, 'Expert panel': 8 board-certified ICU specialists, contribution range 48-123 ratings"),
    ("7", "Variables", "Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers. Give diagnostic criteria, if applicable",
     "Methods, 'Evaluation framework': 10 dimensions defined with descriptions; vote categories defined"),
    ("8", "Data sources/measurement", "For each variable of interest, give sources of data and details of methods of assessment (measurement). Describe comparability of assessment methods if there is more than one group",
     "Methods, 'Evaluation framework': 5-point Likert scale; 'LLM response generation': base Meditron-3 (70B) without fine-tuning"),
    ("9", "Bias", "Describe any efforts to address potential sources of bias",
     "Methods, 'Statistical analysis': simulation null test for circularity; clustered bootstrap for rater dependency. Limitations section addresses selection bias (Vote=12 design), incomplete block design, and lack of rater calibration"),
    ("10", "Study size", "Explain how the study size was arrived at",
     "Methods, 'Expert panel' and 'Question generation': 200 questions by 8 specialists; 658 ratings yielding 788 answer evaluations. Study size determined by MOOVE-CHUV parent study design"),
    ("11", "Quantitative variables", "Explain how quantitative variables were handled in the analyses. If applicable, describe which groupings were chosen and why",
     "Methods, 'Statistical analysis': ordinal weights for vote scale; agreement quartiles based on mean SD; Likert 1-5 scale; BH correction for multiple testing"),
    ("12", "Statistical methods", "(a) Describe all statistical methods, including those used to control for confounding",
     "Methods, 'Statistical analysis': Fleiss' kappa, Krippendorff's alpha, Gwet's AC; Mann-Whitney U with Cliff's delta; Pearson/Spearman correlations; BH correction; clustered bootstrap"),
    ("", "", "(b) Describe any methods used to examine subgroups and interactions",
     "Methods, 'Statistical analysis': stratification by agreement quartiles; subgroup analyses by subspecialty and task type (stated as exploratory)"),
    ("", "", "(c) Explain how missing data were addressed",
     "Limitations: incomplete block design acknowledged; second answer evaluations only available when Vote=12; Krippendorff's alpha selected for tolerance of missing raters"),
    ("", "", "(d) If applicable, describe analytical methods taking account of sampling strategy",
     "Not applicable (census of ICU subcohort within MOOVE-CHUV)"),
    ("", "", "(e) Describe any sensitivity analyses",
     "Supplement: clustered bootstrap CIs; simulation null test; nominal vs ordinal weights"),

    # Results
    ("13", "Participants", "(a) Report numbers of individuals at each stage of study — e.g. numbers potentially eligible, examined for eligibility, confirmed eligible, included in the study, completing follow-up, and analysed",
     "Results: 658 ratings from 8 raters, 788 answer evaluations from 200 questions; rater distribution in Table S1"),
    ("", "", "(b) Give reasons for non-participation at each stage",
     "Methods, 'Expert panel': incomplete block design; median 3 raters per question (range 2-4)"),
    ("", "", "(c) Consider use of a flow diagram",
     "Not included — consider adding a participant flow diagram showing 200 questions -> 658 ratings -> 788 evaluations"),
    ("14", "Descriptive data", "(a) Give characteristics of study participants (e.g. demographic, clinical, social) and information on exposures and potential confounders",
     "Methods, 'Expert panel': board-certified ICU specialists; rater contribution range. Table S1 (rater distribution)"),
    ("", "", "(b) Indicate number of participants with missing data for each variable of interest",
     "Data Dictionary: non-null counts per column; second answer evaluations available for 130/658 ratings (Vote=12 only)"),
    ("15", "Outcome data", "Report numbers of outcome events or summary measures",
     "Table 1: mean, SD, median, IQR, min, max for all 10 dimensions (n=788). Table 2: vote distribution and agreement metrics"),
    ("16", "Main results", "(a) Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision (e.g., 95% confidence interval). Make clear which confounders were adjusted for and why they were included",
     "Tables 1-5: means with SDs; agreement coefficients with 95% CIs; correlations with CIs (Fisher and clustered bootstrap)"),
    ("", "", "(b) Report category boundaries when continuous variables were categorized",
     "Methods/Results: agreement quartiles (Q1-Q4) based on mean SD; alignment bins (1-2, 2-3, 3-4, 4-5)"),
    ("", "", "(c) If relevant, consider translating estimates of relative risk into absolute risk for a meaningful time period",
     "Not applicable"),
    ("17", "Other analyses", "Report other analyses done — e.g. analyses of subgroups and interactions, and sensitivity analyses",
     "Results: stratified analysis (Table 3); correlation analysis; subspecialty analysis (Table 4); task type analysis (Table 5). Supplement: simulation null test (Figure S1); clustered bootstrap; cross-specialty comparison (Table S3); tail risk analysis (Table S4)"),

    # Discussion
    ("18", "Key results", "Summarise key results with reference to study objectives",
     "Discussion, 'Principal findings': performance pattern, agreement levels, alignment-agreement association"),
    ("19", "Limitations", "Discuss limitations of the study, taking into account sources of potential bias or imprecision. Discuss both direction and magnitude of any potential bias",
     "Limitations paragraph: single-center, unequal rater contributions, unvalidated classification, base model only, small subgroups, no rater calibration, meta-evaluation dimensions, Vote=12 selection"),
    ("20", "Interpretation", "Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence",
     "Discussion: comparison with nuclear medicine evaluation; comparison with other LLM evaluation studies; clinical implications stated cautiously ('not ready for autonomous clinical deployment')"),
    ("21", "Generalisability", "Discuss the generalisability (external validity) of the study results",
     "Discussion, 'Comparison with existing literature': consistency with nuclear medicine evaluation suggests intrinsic model properties; 'Base model versus aligned model': results specific to base model"),

    # Other information
    ("22", "Funding", "Give the source of funding and the role of the funders for the present study and, if applicable, for the original study on which the present article is based",
     "Declarations: [TBD]"),
]


def create_strobe_checklist():
    """Create STROBE checklist as a Word document."""

    doc = Document()

    # Title
    title = doc.add_heading('STROBE Statement — Checklist of Items for Cross-Sectional Studies', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Study info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Expert Evaluation of Meditron-3 Large Language Model in Intensive Care Medicine:\n'
        'A Multi-rater Assessment Study'
    )
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Item No.', 'Section', 'Recommendation', 'Reported on page / section']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        # Shade header
        from docx.oxml.ns import qn
        shading = header_cells[i]._element.get_or_add_tcPr()
        shading_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9E2F3',
            qn('w:val'): 'clear',
        })
        shading.append(shading_elem)

    # Data rows
    for item_no, section, recommendation, location in STROBE_ITEMS:
        row = table.add_row()
        cells = row.cells

        cells[0].text = item_no
        cells[1].text = section
        cells[2].text = recommendation
        cells[3].text = location

        # Format cells
        for j, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    if j == 1:
                        run.bold = True if item_no else False

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.0)
        row.cells[2].width = Cm(8.0)
        row.cells[3].width = Cm(6.0)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer.add_run(
        'STROBE = Strengthening the Reporting of Observational Studies in Epidemiology.\n'
        'Reference: von Elm E, Altman DG, Egger M, Pocock SJ, Gotzsche PC, '
        'Vandenbroucke JP. The Strengthening the Reporting of Observational Studies '
        'in Epidemiology (STROBE) statement: guidelines for reporting observational '
        'studies. Lancet. 2007;370(9596):1453-1457.'
    )
    run.font.size = Pt(8)
    run.italic = True

    # Note about item 13c
    note = doc.add_paragraph()
    run = note.add_run(
        'Note: Item 13(c) — a flow diagram showing 200 questions -> 658 ratings -> '
        '788 answer evaluations may be considered for the final submission.'
    )
    run.font.size = Pt(8)
    run.italic = True

    # Save
    output_path = OUTPUT_DIR / 'STROBE_checklist.docx'
    doc.save(output_path)
    return output_path


def main():
    print("Generating STROBE checklist...")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = create_strobe_checklist()
    print(f"Saved to: {path}")


if __name__ == '__main__':
    main()
